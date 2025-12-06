from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.http import HttpResponse
from django.utils import timezone
from datetime import timedelta
from .models import Document
from .serializers import DocumentSerializer, DocumentCreateSerializer, DocumentUpdateSerializer
from analytics.models import AuditLog, UsageMetric
import os
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import io


class DocumentViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = Document.objects.filter(user=self.request.user)
        
        # Filter by priority
        priority = self.request.query_params.get('priority')
        if priority:
            queryset = queryset.filter(priority=priority)
        
        # Filter by status
        doc_status = self.request.query_params.get('status')
        if doc_status:
            queryset = queryset.filter(status=doc_status)
        
        # Filter by file type
        file_type = self.request.query_params.get('file_type')
        if file_type:
            queryset = queryset.filter(file_type=file_type)
        
        # Filter by AI generated
        ai_generated = self.request.query_params.get('ai_generated')
        if ai_generated is not None:
            queryset = queryset.filter(is_ai_generated=ai_generated == 'true')
        
        # Search by title
        search = self.request.query_params.get('search')
        if search:
            queryset = queryset.filter(title__icontains=search)
        
        return queryset
    
    def get_serializer_class(self):
        if self.action == 'create':
            return DocumentCreateSerializer
        if self.action in ['update', 'partial_update']:
            return DocumentUpdateSerializer
        return DocumentSerializer
    
    def perform_create(self, serializer):
        file = serializer.validated_data.get('file')
        
        if file:
            document = serializer.save(
                user=self.request.user,
                original_filename=file.name,
                mime_type=file.content_type,
                file_size=file.size
            )
            
            # Extract text content
            try:
                file.seek(0)  # Reset file pointer
                if file.name.endswith('.pdf'):
                    reader = PdfReader(file)
                    text = '\n'.join([page.extract_text() for page in reader.pages if page.extract_text()])
                    document.content_text = text
                elif file.name.endswith('.docx'):
                    doc = DocxDocument(file)
                    text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                    document.content_text = text
                file.seek(0)  # Reset again for storage
            except Exception as e:
                pass  # Handle errors silently for MVP
            
            document.save()
        else:
            # For AI-generated documents without file
            document = serializer.save(
                user=self.request.user,
                is_ai_generated=True
            )
        
        # Track usage
        UsageMetric.objects.create(
            user=self.request.user,
            metric_type='document_uploaded',
            value=1,
            metadata={'document_id': document.id}
        )
        
        # Audit log
        AuditLog.objects.create(
            user=self.request.user,
            action='document_access',
            resource_type='document',
            resource_id=document.id
        )
    
    @action(detail=True, methods=['post'])
    def analyze(self, request, pk=None):
        """Analyze document with AI"""
        document = self.get_object()
        if not document.content_text:
            return Response(
                {'error': 'Document has no extractable text'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        from ai_agent.services import AIService
        ai_service = AIService()
        
        try:
            result = ai_service.analyze_document(
                document.content_text,
                language=request.user.language
            )
            
            # Save analysis to document
            document.analysis = result['content']
            document.save()
            
            UsageMetric.objects.create(
                user=request.user,
                metric_type='document_analyzed',
                value=result['tokens_used'],
                metadata={'document_id': document.id}
            )
            
            return Response({
                'analysis': result['content'],
                'document_id': document.id
            })
        except Exception as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """Download document file"""
        document = self.get_object()
        
        if not document.file:
            return Response(
                {'error': 'Document has no file attached'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        AuditLog.objects.create(
            user=request.user,
            action='document_access',
            resource_type='document',
            resource_id=document.id,
            metadata={'action': 'download'}
        )
        
        response = HttpResponse(document.file.read(), content_type=document.mime_type or 'application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{document.original_filename or document.title}"'
        return response
    
    @action(detail=True, methods=['get', 'post'])
    def export_to_docx(self, request, pk=None):
        """Export document content to DOCX format"""
        document = self.get_object()
        
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches
            
            doc = DocxDocument()
            doc.add_heading(document.title, 0)
            
            content = document.content_text or document.analysis or ''
            if content:
                for paragraph in content.split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
            
            # Save to memory
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            AuditLog.objects.create(
                user=request.user,
                action='document_access',
                resource_type='document',
                resource_id=document.id,
                metadata={'action': 'export_docx'}
            )
            
            response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{document.title}.docx"'
            return response
        except Exception as e:
            return Response(
                {'error': f'Export failed: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['patch'])
    def update_priority(self, request, pk=None):
        """Update document priority"""
        document = self.get_object()
        priority = request.data.get('priority')
        
        if priority not in ['low', 'medium', 'urgent']:
            return Response(
                {'error': 'Invalid priority. Must be low, medium, or urgent'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        document.priority = priority
        document.save()
        
        return Response(DocumentSerializer(document).data)
    
    @action(detail=True, methods=['patch'])
    def update_status(self, request, pk=None):
        """Update document status"""
        document = self.get_object()
        doc_status = request.data.get('status')
        
        if doc_status not in ['started', 'in-progress', 'done']:
            return Response(
                {'error': 'Invalid status. Must be started, in-progress, or done'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        document.status = doc_status
        document.save()
        
        return Response(DocumentSerializer(document).data)
    
    @action(detail=True, methods=['patch'])
    def update_tags(self, request, pk=None):
        """Update document tags"""
        document = self.get_object()
        tags = request.data.get('tags', [])
        
        if not isinstance(tags, list):
            return Response(
                {'error': 'Tags must be a list'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        document.tags = tags
        document.save()
        
        return Response(DocumentSerializer(document).data)
    
    def destroy(self, request, *args, **kwargs):
        document = self.get_object()
        
        # Audit log
        AuditLog.objects.create(
            user=request.user,
            action='document_delete',
            resource_type='document',
            resource_id=document.id
        )
        
        return super().destroy(request, *args, **kwargs)
