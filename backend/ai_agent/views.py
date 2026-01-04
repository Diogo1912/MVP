from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated, AllowAny
from django.utils import timezone
from django.conf import settings
from .models import Conversation, Message, Prompt, KnowledgeBase
from .services import AIService
from documents.models import Document
from analytics.models import UsageMetric
import json
import sys
import logging

logger = logging.getLogger(__name__)


class ConversationViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Conversation.objects.filter(user=self.request.user).order_by('-updated_at')
    
    def get_serializer_class(self):
        from .serializers import ConversationSerializer
        return ConversationSerializer
    
    def perform_create(self, serializer):
        serializer.save(user=self.request.user, language=self.request.user.language)
    
    @action(detail=True, methods=['get'])
    def messages(self, request, pk=None):
        """Get all messages for a conversation"""
        conversation = self.get_object()
        messages = Message.objects.filter(conversation=conversation).order_by('created_at')
        from .serializers import MessageSerializer
        return Response(MessageSerializer(messages, many=True).data)


class MessageViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        conversation_id = self.request.query_params.get('conversation')
        if conversation_id:
            return Message.objects.filter(
                conversation_id=conversation_id,
                conversation__user=self.request.user
            ).order_by('created_at')
        return Message.objects.filter(conversation__user=self.request.user).order_by('-created_at')
    
    def get_serializer_class(self):
        from .serializers import MessageSerializer
        return MessageSerializer


class PromptViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Prompt.objects.all().order_by('name', '-version')
    
    def get_serializer_class(self):
        from .serializers import PromptSerializer
        return PromptSerializer
    
    @action(detail=False, methods=['get'])
    def active(self, request):
        """Get all active prompts"""
        prompts = Prompt.objects.filter(is_active=True)
        from .serializers import PromptSerializer
        return Response(PromptSerializer(prompts, many=True).data)


class KnowledgeBaseViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return KnowledgeBase.objects.filter(document__user=self.request.user)
    
    def get_serializer_class(self):
        from .serializers import KnowledgeBaseSerializer
        return KnowledgeBaseSerializer


class ChatView(APIView):
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        """Handle chat message and get AI response"""
        user = request.user
        message_content = request.data.get('message', '')
        conversation_id = request.data.get('conversation_id')
        document_id = request.data.get('document_id')
        persona = request.data.get('persona', 'commercial')  # commercial or personal
        use_knowledge_base = request.data.get('use_knowledge_base', False)
        case_id = request.data.get('case_id')  # Case to assign or update
        
        if not message_content:
            return Response({'error': 'Message is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get or create conversation
        if conversation_id:
            try:
                conversation = Conversation.objects.get(id=conversation_id, user=user)
                # Update case if provided
                if case_id:
                    conversation.case_id = case_id
                    conversation.save()
            except Conversation.DoesNotExist:
                return Response({'error': 'Conversation not found'}, status=status.HTTP_404_NOT_FOUND)
        else:
            conversation = Conversation.objects.create(
                user=user,
                language=user.language,
                title=message_content[:50],
                case_id=case_id if case_id else None
            )
        
        # Get document context if provided
        document_context = None
        document = None
        if document_id:
            try:
                document = Document.objects.get(id=document_id, user=user)
                document_context = document.content_text[:2000] if document.content_text else None
            except Document.DoesNotExist:
                pass
        
        # Get case context if assigned to conversation
        case_context = None
        if conversation.case_id:
            try:
                from cases.models import Case
                case = Case.objects.get(id=conversation.case_id, lawyer=user)
                # Fetch all documents related to this case
                case_documents = Document.objects.filter(case=case, user=user)
                
                context_parts = [f"Case Title: {case.title}", f"Case Description: {case.description}"]
                for doc in case_documents:
                    if doc.content_text:
                        context_parts.append(f"Document: {doc.title}\n{doc.content_text[:800]}")
                case_context = "\n\n".join(context_parts)
            except:
                pass
        
        # Save user message
        user_message = Message.objects.create(
            conversation=conversation,
            role='user',
            content=message_content,
            document=document
        )
        
        # Get conversation history
        previous_messages = Message.objects.filter(
            conversation=conversation
        ).exclude(id=user_message.id).order_by('created_at')[:10]
        
        messages = [
            {"role": msg.role, "content": msg.content}
            for msg in previous_messages
        ]
        messages.append({"role": "user", "content": message_content})
        
        # Get AI response
        try:
            logger.info(f"Initializing AI service for user {user.email}")
            ai_service = AIService()
            logger.info("AI service initialized successfully")
        except Exception as e:
            # AI service initialization failed
            import traceback
            error_msg = str(e)
            logger.error(f"AI Service Init Error: {error_msg}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return Response(
                {'error': f"AI service initialization failed: {error_msg}"},
                status=status.HTTP_503_SERVICE_UNAVAILABLE
            )
        
        try:
            logger.info(f"Attempting chat completion for user {user.email}, persona: {persona}")
            response = ai_service.chat_completion(
                messages=messages,
                language=user.language,
                persona=persona,
                use_knowledge_base=use_knowledge_base,
                document_context=document_context,
                case_context=case_context
            )
            logger.info(f"Chat completion successful, tokens used: {response.get('tokens_used', 0)}")
            
            # Save AI message
            ai_message = Message.objects.create(
                conversation=conversation,
                role='assistant',
                content=response['content'],
                tokens_used=response['tokens_used'],
                document=document
            )
            
            # Update conversation
            conversation.updated_at = timezone.now()
            conversation.save()
            
            # Track usage
            UsageMetric.objects.create(
                user=user,
                metric_type='ai_query',
                value=response['tokens_used'],
                metadata={
                    'conversation_id': conversation.id,
                    'persona': persona
                }
            )
            
            from .serializers import MessageSerializer
            return Response({
                'conversation_id': conversation.id,
                'message': MessageSerializer(ai_message).data,
                'user_message': MessageSerializer(user_message).data,
            })
            
        except Exception as e:
            import traceback
            error_msg = str(e)
            logger.error(f"AI Chat Completion Error: {error_msg}")
            logger.error(f"Full Traceback: {traceback.format_exc()}")
            logger.error(f"Error type: {type(e).__name__}")
            
            # Return detailed error for debugging
            return Response(
                {
                    'error': error_msg,
                    'error_type': type(e).__name__,
                    'details': 'Check deploy logs for full traceback'
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class RegenerateView(APIView):
    """Regenerate an AI response with additional instructions"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        message_id = request.data.get('message_id')
        additional_instructions = request.data.get('instructions', '')
        
        if not message_id:
            return Response({'error': 'Message ID is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Get the original AI message
            ai_message = Message.objects.get(
                id=message_id,
                role='assistant',
                conversation__user=request.user
            )
            
            # Get the user message that triggered this response
            user_message = Message.objects.filter(
                conversation=ai_message.conversation,
                role='user',
                created_at__lt=ai_message.created_at
            ).order_by('-created_at').first()
            
            if not user_message:
                return Response({'error': 'Original user message not found'}, status=status.HTTP_404_NOT_FOUND)
            
            # Regenerate response
            ai_service = AIService()
            response = ai_service.regenerate_response(
                original_message=user_message.content,
                previous_response=ai_message.content,
                additional_instructions=additional_instructions,
                language=request.user.language
            )
            
            # Create new AI message
            new_ai_message = Message.objects.create(
                conversation=ai_message.conversation,
                role='assistant',
                content=response['content'],
                tokens_used=response['tokens_used'],
                document=ai_message.document
            )
            
            # Track usage
            UsageMetric.objects.create(
                user=request.user,
                metric_type='ai_query',
                value=response['tokens_used'],
                metadata={
                    'conversation_id': ai_message.conversation.id,
                    'regeneration': True
                }
            )
            
            from .serializers import MessageSerializer
            return Response({
                'message': MessageSerializer(new_ai_message).data,
            })
            
        except Message.DoesNotExist:
            return Response({'error': 'Message not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class GenerateDocumentView(APIView):
    """Generate a document from AI content and create actual DOCX file"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        content = request.data.get('content', '')
        title = request.data.get('title', 'AI Generated Document')
        document_type = request.data.get('type', 'docx')
        case_id = request.data.get('case_id')
        
        if not content:
            return Response({'error': 'Content is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from django.core.files.base import ContentFile
            import io
            import re
            
            # Create a professional DOCX document
            doc = DocxDocument()
            
            # Set document title
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date
            from datetime import datetime
            date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()  # Spacer
            
            # Process content - handle markdown-like formatting
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                
                # Handle headers
                if line.startswith('###'):
                    doc.add_heading(line.replace('###', '').strip(), level=3)
                elif line.startswith('##'):
                    doc.add_heading(line.replace('##', '').strip(), level=2)
                elif line.startswith('#'):
                    doc.add_heading(line.replace('#', '').strip(), level=1)
                elif line.startswith('**') and line.endswith('**'):
                    # Bold section header
                    para = doc.add_paragraph()
                    run = para.add_run(line.replace('**', ''))
                    run.bold = True
                elif line.startswith('- ') or line.startswith('• '):
                    # Bullet point
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\.', line):
                    # Numbered list
                    doc.add_paragraph(line, style='List Number')
                else:
                    # Regular paragraph
                    para = doc.add_paragraph()
                    # Handle inline bold
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = para.add_run(part[2:-2])
                            run.bold = True
                        else:
                            para.add_run(part)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Create filename
            safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:50]
            filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            # Create document record with actual file
            document = Document.objects.create(
                user=request.user,
                title=title,
                file_type='ai_generated',
                content_text=content,
                is_ai_generated=True,
                original_filename=filename,
                mime_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                case_id=case_id if case_id else None
            )
            
            # Save the DOCX file
            document.file.save(filename, ContentFile(buffer.getvalue()))
            document.file_size = len(buffer.getvalue())
            document.save()
            
            # Track usage
            UsageMetric.objects.create(
                user=request.user,
                metric_type='document_created',
                value=1,
                metadata={'document_id': document.id, 'ai_generated': True}
            )
            
            from documents.serializers import DocumentSerializer
            return Response({
                'document': DocumentSerializer(document).data,
            })
            
        except Exception as e:
            logger.error(f"Document generation error: {str(e)}")
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class AIHealthCheckView(APIView):
    """Health check endpoint for AI service"""
    permission_classes = [AllowAny]
    
    def get(self, request):
        health_status = {
            'openai_installed': False,
            'openai_key_set': False,
            'ai_service_init': False,
            'python_version': sys.version,
            'errors': []
        }
        
        # Check if openai is installed
        try:
            import openai
            health_status['openai_installed'] = True
            health_status['openai_version'] = openai.__version__
        except ImportError as e:
            health_status['errors'].append(f'OpenAI not installed: {str(e)}')
        
        # Check if API key is set
        api_key = getattr(settings, 'OPENAI_API_KEY', None)
        if api_key:
            health_status['openai_key_set'] = True
            health_status['api_key_length'] = len(api_key)
        else:
            health_status['errors'].append('OPENAI_API_KEY not set in environment')
        
        # Try to initialize AI service
        try:
            ai_service = AIService()
            health_status['ai_service_init'] = True
        except Exception as e:
            health_status['errors'].append(f'AI Service init failed: {str(e)}')
        
        # Return status
        status_code = status.HTTP_200_OK if not health_status['errors'] else status.HTTP_503_SERVICE_UNAVAILABLE
        return Response(health_status, status=status_code)
